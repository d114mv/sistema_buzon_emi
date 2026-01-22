import csv
import os
from django.conf import settings
from django.contrib import admin, messages
from django.urls import reverse
from django.utils.html import format_html
from django.http import HttpResponse

from docx import Document 
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH

from .models import Categoria, Ticket

admin.site.site_header = "Panel de Control EMI"
admin.site.site_title = "Buzón EMI"
admin.site.index_title = "Gestión de Reportes y Sugerencias"

@admin.register(Ticket)
class TicketAdmin(admin.ModelAdmin):
    list_display = ('miniatura', 'asunto_corto', 'categoria', 'status_coloreado', 'fecha_creacion', 'ver_detalle_boton')
    list_filter = ('estado', 'categoria', 'fecha_creacion')
    search_fields = ('asunto', 'descripcion')
    readonly_fields = ('fecha_creacion', 'usuario_hash', 'vista_previa_grande')

    actions = ['generar_informe_word', 'exportar_a_csv', 'marcar_resuelto', 'marcar_proceso']

    fieldsets = (
        ('Información del Reporte', {
            'fields': ('categoria', 'asunto', 'descripcion', 'fecha_creacion')
        }),
        ('Evidencia', {
            'fields': ('imagen', 'vista_previa_grande')
        }),
        ('Gestión Administrativa', {
            'fields': ('estado', 'comentario_admin', 'usuario_hash')
        }),
    )

    @admin.action(description="📝 Generar Informe Oficial (Word)")
    def generar_informe_word(self, request, queryset):
        if queryset.count() > 1:
            self.message_user(request, "⚠️ Por favor, selecciona SOLO UN ticket para el informe Word.", level=messages.ERROR)
            return
        
        ticket = queryset.first()
        doc = Document()
        
        titulo = doc.add_heading('ESCUELA MILITAR DE INGENIERÍA', 0)
        titulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        subtitulo = doc.add_paragraph('REPORTE DE INCIDENTE / SUGERENCIA')
        subtitulo.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph('_' * 70)

        table = doc.add_table(rows=6, cols=2)
        table.style = 'Table Grid'

        def fill_row(row_idx, label, value):
            row = table.rows[row_idx]
            row.cells[0].text = label
            row.cells[0].paragraphs[0].runs[0].font.bold = True
            row.cells[1].text = str(value)

        fill_row(0, 'Código de Reporte:', ticket.id)
        fill_row(1, 'Fecha:', ticket.fecha_creacion.strftime("%d/%m/%Y %H:%M"))
        fill_row(2, 'Categoría:', ticket.categoria.nombre)
        fill_row(3, 'Prioridad:', f"Nivel {ticket.categoria.prioridad_base}")
        fill_row(4, 'Asunto:', ticket.asunto)
        fill_row(5, 'Estado:', ticket.get_estado_display())

        doc.add_heading('Detalle:', level=2)
        p = doc.add_paragraph(ticket.descripcion)
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY

        doc.add_heading('Evidencia:', level=2)
        if ticket.imagen:
            try:
                image_path = os.path.join(settings.MEDIA_ROOT, ticket.imagen.name)
                doc.add_picture(image_path, width=Inches(4))
                doc.add_paragraph('(Imagen adjunta)')
            except Exception as e:
                doc.add_paragraph(f"[Error cargando imagen: {e}]")
        else:
            doc.add_paragraph('Sin evidencia fotográfica.')

        linea = doc.add_paragraph('\n\n\n' + ('_' * 40))
        linea.alignment = WD_ALIGN_PARAGRAPH.CENTER
        firma = doc.add_paragraph('Firma y Sello del Responsable')
        firma.alignment = WD_ALIGN_PARAGRAPH.CENTER

        response = HttpResponse(content_type='application/vnd.openxmlformats-officedocument.wordprocessingml.document')
        filename = f"Informe_EMI_{str(ticket.id)[:6]}.docx"
        response['Content-Disposition'] = f'attachment; filename="{filename}"'
        doc.save(response)
        return response

    @admin.action(description="📄 Descargar Reporte Masivo (Excel/CSV)")
    def exportar_a_csv(self, request, queryset):
        response = HttpResponse(content_type='text/csv')
        response['Content-Disposition'] = 'attachment; filename="reporte_quejas_emi.csv"'
        writer = csv.writer(response)
        writer.writerow(['ID', 'Fecha', 'Categoria', 'Asunto', 'Descripcion', 'Estado', 'Tiene Foto'])
        for ticket in queryset:
            tiene_foto = "SI" if ticket.imagen else "NO"
            writer.writerow([
                ticket.id, ticket.fecha_creacion.strftime("%Y-%m-%d %H:%M"),
                ticket.categoria.nombre, ticket.asunto, ticket.descripcion,
                ticket.get_estado_display(), tiene_foto
            ])
        return response

    @admin.action(description="Marcar como RESUELTO")
    def marcar_resuelto(self, request, queryset):
        queryset.update(estado='RES')

    @admin.action(description="Marcar como EN PROCESO")
    def marcar_proceso(self, request, queryset):
        queryset.update(estado='PROC')

    def ver_detalle_boton(self, obj):
        url = reverse('admin:tickets_ticket_change', args=[obj.id])
        return format_html('<a class="button" style="background-color:#003366; color:white; padding:5px 10px; border-radius:5px; text-decoration:none;" href="{}">Ver Informe</a>', url)
    ver_detalle_boton.short_description = "Acción"

    def miniatura(self, obj):
        if obj.imagen:
            return format_html('<img src="{}" style="width: 50px; height:50px; object-fit:cover; border-radius:4px;" />', obj.imagen.url)
        return "❌"
    miniatura.short_description = "Foto"

    def vista_previa_grande(self, obj):
        if obj.imagen:
            return format_html('<img src="{}" style="max-width: 100%; height:auto; border-radius:8px;" />', obj.imagen.url)
        return "Sin evidencia"
    vista_previa_grande.short_description = "Vista Previa"

    def asunto_corto(self, obj):
        return obj.asunto[:30] + "..." if len(obj.asunto) > 30 else obj.asunto
    asunto_corto.short_description = "Asunto"

    def status_coloreado(self, obj):
        colores = {'PEND': 'red', 'PROC': 'orange', 'RES': 'green', 'RECH': 'gray'}
        return format_html('<strong style="color: {};">{}</strong>', colores.get(obj.estado, 'black'), obj.get_estado_display())
    status_coloreado.short_description = "Estado"

@admin.register(Categoria)
class CategoriaAdmin(admin.ModelAdmin):
    list_display = ('nombre', 'email_responsable', 'prioridad_base')